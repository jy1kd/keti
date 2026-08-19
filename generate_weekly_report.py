# -*- coding: utf-8 -*-
"""
Generate weekly internship report for 2026-07-27 to 2026-07-31
based on the template "实习周报.docx" in the keti directory.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_text(cell, text, bold=False, size=12, align="center"):
    """Clear cell paragraphs and set new text."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if align == "center" else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold

def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(title)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.bold = True

def add_paragraph(doc, text, indent=True, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.bold = bold

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.first_line_indent = Pt(24 if level == 0 else 48)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in p.runs:
                r.font.name = "宋体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(11)
                r.font.bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for r in p.runs:
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    r.font.size = Pt(10.5)
    return table


def main():
    doc = Document("实习周报.docx")
    table = doc.tables[0]

    # Fill header cells
    set_cell_text(table.rows[0].cells[1], "原毅博", size=12)
    set_cell_text(table.rows[0].cells[3], "2026.7.31", size=12)
    set_cell_text(table.rows[1].cells[1], "现货团队", size=12)
    set_cell_text(table.rows[1].cells[3], "郑昊泽", size=12)
    set_cell_text(table.rows[2].cells[1], "2026年7月27日-2026年7月31日", size=12)

    # Clear the big merged content cell and use it as the document body
    content_cell = table.rows[3].cells[0]
    content_cell.text = ""

    # We will add content paragraphs directly to the cell
    # Title
    p = content_cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("一、学习心得")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.bold = True

    add_paragraph_to_cell(content_cell,
        "1. Electron 桌面应用架构实践\n"
        "本周系统学习了 Electron 主进程/渲染进程分离模型，掌握了 IPC 通信、BrowserWindow 多窗口管理、"
        "系统托盘（Tray）、全局快捷键（globalShortcut）及原生通知（Notification）等核心模块。"
        "通过将现有 React + Vite Web 应用迁移为桌面应用，理解了如何在不破坏前端业务代码的前提下，"
        "通过预加载脚本（preload）安全暴露主进程能力。")

    add_paragraph_to_cell(content_cell,
        "2. 桌面应用打包与后端集成\n"
        "学习了 electron-builder 跨平台打包配置、NSIS 安装程序、自动更新（electron-updater）机制；"
        "同时了解了如何使用 PyInstaller 将 Python FastAPI 后端打包为独立可执行文件，"
        "并通过 BackendManager 在 Electron 应用启动/关闭时自动管理后端进程生命周期。")

    add_paragraph_to_cell(content_cell,
        "3. 大型前端重构的方案设计\n"
        "参与了多页面架构 + 虚拟滚动按需订阅的重构方案设计，学习了如何用标签页系统替代现有三栏布局，"
        "以及如何通过可见行检测、防抖批量订阅/退订来降低行情数据压力，支撑全量合约（6000+）展示。")

    # Section 2
    p = content_cell.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("二、工作结果")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.bold = True

    add_paragraph_to_cell(content_cell,
        "1. Electron 桌面应用迁移完成（PR-E1 ~ PR-E10，PR #57）\n"
        "本周完成了从 Web 应用到 Electron 桌面应用的全套迁移，共 10 个 PR，覆盖框架搭建、多窗口、托盘、快捷键、通知、打包及后端集成。")

    # Add PR table
    pr_table = content_cell.add_table(rows=1, cols=5)
    pr_table.style = "Table Grid"
    hdr = pr_table.rows[0].cells
    headers = ["PR", "标题", "完成时间", "关键文件", "状态"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        fmt_cell(hdr[i], bold=True)
    pr_data = [
        ["PR-E1", "Electron 基础框架搭建", "7.28", "electron/main.ts, preload.ts", "✅"],
        ["PR-E2", "IPC 通信基础设施", "7.28", "electron/ipc/, src/services/electron.ts", "✅"],
        ["PR-E3", "窗口管理器实现", "7.28", "electron/windowManager.ts", "✅"],
        ["PR-E4", "报单窗口实现", "7.30", "electron/windows/orderWindow.ts", "✅"],
        ["PR-E5", "K 线窗口实现", "7.30", "electron/windows/klineWindow.ts", "✅"],
        ["PR-E6", "系统托盘实现", "7.31", "electron/trayManager.ts", "✅"],
        ["PR-E7", "全局快捷键实现", "7.31", "electron/shortcuts.ts", "✅"],
        ["PR-E8", "原生通知实现", "7.31", "electron/notificationManager.ts", "✅"],
        ["PR-E9", "应用打包配置", "7.31", "electron-builder.json, build-electron.js", "✅"],
        ["PR-E10", "Python 后端打包集成", "7.31", "electron/backendManager.ts, pyinstaller.spec", "✅"],
    ]
    for row in pr_data:
        cells = pr_table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            fmt_cell(cells[i])

    add_paragraph_to_cell(content_cell,
        "关键修复：修复 ES Module 兼容性问题并自动将 dist-electron 中的 .js 重命名为 .cjs；"
        "修复 trayManager.ts / autoUpdater.ts TypeScript 类型错误；"
        "修复 BackendManager 检测已有后端进程逻辑；清理构建产物仅保留 .cjs。")

    add_paragraph_to_cell(content_cell,
        "2. K 线功能完善与技术指标集成（7.30）\n"
        "实现技术指标计算函数（成交量均线、布林带、KDJ、RSI），集成主图指标（MA5/MA10/MA20、BOLL）"
        "和副图指标（成交量/VOL-MA5、MACD、KDJ、RSI）。修复 K 线时间戳统一、成交量虚高、high/low 周期计算、夜盘时间戳等问题，删除日线功能。")

    add_paragraph_to_cell(content_cell,
        "3. 一致性检查与 Bug 修复（7.27-7.28）\n"
        "完成一致性检查 check04/check05/check06，修复类型定义、文档对齐、optionType 命名统一等问题。"
        "修复前后端代码审查发现的 21 个 Bug；修复交易指令合规性（保护价/数量上限/涨跌停/止损市价/exchangeID）；"
        "修复一键反向/锁仓保护价为 0、止损单市价触发保护价传递、撤单 sessionID、部分成交状态更新等关键缺陷。")

    add_paragraph_to_cell(content_cell,
        "4. 文档重构与下阶段设计（7.28-7.31）\n"
        "重组 docs 目录结构并统一英文命名；更新 README.md、CLAUDE.md，补充 Electron 架构图和桌面应用文档；"
        "添加 task-electron-migration.md、task-redesign.md；设计标签页系统、收藏功能、右键菜单、IPC 监控、设置面板等。")

    add_paragraph_to_cell(content_cell,
        "5. 本周关键数据\n"
        "合并 PR #40~#45、#52~#57 等约 18 个分支；7.27-7.31 累计新增/修复 commit 约 200 个；"
        "后端测试保持 391+ tests passed。当前阶段：Web 功能基本闭环，Electron 桌面应用迁移完成，进入下一轮 UI/UX 重构设计阶段。")

    # Section 3
    p = content_cell.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("三、下周计划")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.font.bold = True

    plans = [
        "启动多页面架构 + 标签页系统重构开发（PR-R1 ~ PR-R5）",
        "完成行情表格虚拟滚动 + 按需订阅实现",
        "实现收藏功能替代原有订阅/退订逻辑",
        "配合角色 B 完成 Electron 与标签页系统的集成（PR-R21/R22）",
        "准备端到端联调测试（PR-17）",
    ]
    for plan in plans:
        p = content_cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-12)
        run = p.add_run("• " + plan)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)

    # Remove extra empty paragraphs at the end of the cell if any
    # Save
    output_path = "实习周报原毅博Week5.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


def add_paragraph_to_cell(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def fmt_cell(cell, bold=False):
    for p in cell.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for r in p.runs:
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(10.5)
            r.font.bold = bold


if __name__ == "__main__":
    main()
